# Fichier : utils/file_utils.py
# CORRECTION (Phase 2) : Les fonctions reçoivent maintenant directement l'instance
# du CongeManager au lieu du db_path. Cela évite de créer des connexions
# inutiles à la base de données et améliore la performance et la fiabilité.

import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font
from datetime import datetime
import re
import logging
import docx

from utils.config_loader import CONFIG
from utils.date_utils import format_date_for_display

def export_agents_to_excel(manager, save_path):
    """Exporte la liste des agents. Conçu pour être exécuté dans un thread."""
    agents = manager.get_all_agents()
    if not agents:
        return "Aucun agent à exporter."
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Agents"
    
    annee_exercice = manager.get_annee_exercice()
    an_n, an_n1, an_n2 = annee_exercice, annee_exercice - 1, annee_exercice - 2
    headers = ["ID", "Nom", "Prénom", "PPR", "Grade", 
               f"Solde {an_n2}", f"Solde {an_n1}", f"Solde {an_n}", "Solde Total Actif"]
    ws.append(headers)
    
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font

    for agent in agents:
        soldes_par_annee = {s.annee: s.solde for s in agent.soldes_annuels if s.statut == 'Actif'}
        solde_n2 = soldes_par_annee.get(an_n2, 0.0)
        solde_n1 = soldes_par_annee.get(an_n1, 0.0)
        solde_n  = soldes_par_annee.get(an_n, 0.0)
        solde_total = agent.get_solde_total_actif()
        
        ws.append([agent.id, agent.nom, agent.prenom, agent.ppr, agent.grade, 
                   solde_n2, solde_n1, solde_n, solde_total])

    for col_idx, col_cells in enumerate(ws.columns, 1):
        max_length = max(len(str(cell.value or "")) for cell in col_cells)
        ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 2
        
    wb.save(save_path)
    return f"Liste des agents exportée avec succès vers\n{save_path}"

def export_all_conges_to_excel(manager, save_path):
    """Exporte la liste de tous les congés. Conçu pour être exécuté dans un thread."""
    all_conges = manager.get_all_conges()
    if not all_conges:
        return "Aucun congé à exporter."
        
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tous les Congés"
    headers = ["Nom Agent", "Prénom Agent", "PPR Agent", "Type Congé", "Début", "Fin", "Jours Pris", "Statut", "Justification", "Intérimaire"]
    ws.append(headers)
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font
        
    all_agents = {agent.id: agent for agent in manager.get_all_agents()}
    for conge in all_conges:
        agent = all_agents.get(conge.agent_id)
        agent_nom, agent_prenom, agent_ppr = (agent.nom, agent.prenom, agent.ppr) if agent else ("Agent", "Supprimé", "")
        interim_info = ""
        if conge.interim_id:
            interim = all_agents.get(conge.interim_id)
            interim_info = f"{interim.nom} {interim.prenom}" if interim else "Agent Supprimé"
        row_data = [agent_nom, agent_prenom, agent_ppr, conge.type_conge, format_date_for_display(conge.date_debut), format_date_for_display(conge.date_fin), conge.jours_pris, conge.statut, conge.justif or "", interim_info]
        ws.append(row_data)
        
    for col_idx, col_cells in enumerate(ws.columns, 1):
        max_length = max(len(str(cell.value or "")) for cell in col_cells)
        ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 2
        
    wb.save(save_path)
    return f"Tous les congés ont été exportés avec succès vers\n{save_path}"

def import_agents_from_excel(manager, source_path):
    """Importe des agents avec une logique de colonnes optionnelles."""
    errors = []
    added_count, updated_count = 0, 0
    
    required_headers = ['nom', 'prenom']
    grades = CONFIG['ui']['grades']
    default_grade = grades[0] if grades else "Administrateur"
    
    wb = openpyxl.load_workbook(source_path)
    ws = wb.active
    
    header = [str(cell.value or '').lower().strip() for cell in ws[1]]
    if not all(h in header for h in required_headers):
        raise ValueError(f"Colonnes requises manquantes : {', '.join(required_headers)}")

    col_map = {name: i for i, name in enumerate(header)}
    
    manager.db.conn.execute('BEGIN TRANSACTION')
    try:
        for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if all(c is None for c in row):
                continue
            try:
                nom = str(row[col_map['nom']] or '').strip()
                prenom = str(row[col_map['prenom']] or '').strip()
                if not nom or not prenom:
                    raise ValueError("Nom et prénom sont obligatoires.")

                ppr = str(row[col_map.get('ppr')] or '').strip() or f"{nom.upper()[:4]}_{datetime.now().strftime('%f')}"
                grade = str(row[col_map.get('grade')] or '').strip() or default_grade
                if grade not in grades:
                    raise ValueError(f"Grade '{grade}' invalide. Grades valides : {', '.join(grades)}")

                soldes = {}
                for col_name, col_idx in col_map.items():
                    match = re.match(r'solde_(\d{4})', col_name)
                    if match and row[col_idx] is not None:
                        annee = int(match.group(1))
                        solde_val = float(str(row[col_idx]).replace(",", "."))
                        if solde_val < 0:
                            raise ValueError(f"Solde négatif pour l'année {annee}.")
                        soldes[annee] = solde_val
                
                agent_data = {
                    'nom': nom, 'prenom': prenom, 'ppr': ppr, 'grade': grade,
                    'soldes': soldes
                }
                
                existing_agent_id = manager.db.execute_query("SELECT id FROM agents WHERE ppr=?", (ppr,), fetch="one")
                if existing_agent_id:
                    agent_data['id'] = existing_agent_id[0]
                    manager.save_agent(agent_data, is_modification=True)
                    updated_count += 1
                else:
                    manager.save_agent(agent_data, is_modification=False)
                    added_count += 1
            
            except Exception as ve:
                logging.warning(f"Erreur d'import à la ligne {i}: {ve}", exc_info=True)
                errors.append(f"Ligne {i}: {ve}")
        
        if errors:
            manager.db.conn.rollback()
            raise Exception("Importation annulée en raison d'erreurs:\n" + "\n".join(errors[:10]))
        else:
            manager.db.conn.commit()
            return f"Importation réussie !\n\n- Agents ajoutés : {added_count}\n- Agents mis à jour : {updated_count}"

    except Exception as e:
        manager.db.conn.rollback()
        raise e

def generate_decision_from_template(template_path, output_path, context):
    """
    Génère un document Word à partir d'un modèle en remplaçant les tags.
    """
    try:
        doc = docx.Document(template_path)

        # Remplacement dans les paragraphes du corps du document
        for p in doc.paragraphs:
            full_text = "".join(run.text for run in p.runs)
            
            if '{{' in full_text and '}}' in full_text:
                for key, value in context.items():
                    full_text = full_text.replace(key, str(value))
                
                for i in range(len(p.runs)):
                    p.runs[i].text = ''
                p.runs[0].text = full_text

        # Remplacement dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full_text = "".join(run.text for run in p.runs)
                        if '{{' in full_text and '}}' in full_text:
                            for key, value in context.items():
                                full_text = full_text.replace(key, str(value))
                            
                            for i in range(len(p.runs)):
                                p.runs[i].text = ''
                            p.runs[0].text = full_text

        doc.save(output_path)
        return True
    except Exception as e:
        logging.error(f"Erreur lors de la génération du document : {e}", exc_info=True)
        raise e